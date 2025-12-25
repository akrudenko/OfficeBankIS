from datetime import datetime
from pathlib import Path
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def save_csv(df: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(path, index=False, encoding="utf-8-sig")

def save_util_chart(util_df: pd.DataFrame, path_png: Path) -> None:
    path_png.parent.mkdir(parents=True, exist_ok=True)
    plt.figure(figsize=(8,4))
    if util_df.empty:
        plt.title("Загрузка: нет данных")
    else:
        plt.plot(util_df["Date"], util_df["UtilizationPct"])
        plt.xticks(rotation=45, ha="right")
        plt.ylabel("Загрузка, %")
        plt.title("Загрузка офиса по дням")
        plt.tight_layout()
    plt.savefig(path_png, dpi=200)
    plt.close()

def save_docx_summary(start_at: datetime, end_at: datetime,
                      util_df: pd.DataFrame, top_df: pd.DataFrame,
                      chart_png: Path, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Отчёт по загрузке офисных ресурсов", level=1)
    doc.add_paragraph(f"Период: {start_at:%d.%m.%Y %H:%M} — {end_at:%d.%m.%Y %H:%M}")

    doc.add_heading("1. Загрузка по дням", level=2)
    if util_df.empty:
        doc.add_paragraph("Данных нет.")
    else:
        t = doc.add_table(rows=1, cols=3)
        t.rows[0].cells[0].text = "Дата"
        t.rows[0].cells[1].text = "Минут занято"
        t.rows[0].cells[2].text = "Загрузка, %"
        for _, r in util_df.iterrows():
            row = t.add_row().cells
            row[0].text = str(r["Date"])
            row[1].text = str(int(r["BookedMinutes"]))
            row[2].text = str(r["UtilizationPct"])

    doc.add_paragraph("")
    if chart_png.exists():
        doc.add_picture(str(chart_png), width=Inches(6))

    doc.add_heading("2. Топ ресурсов", level=2)
    if top_df.empty:
        doc.add_paragraph("Нет данных.")
    else:
        t2 = doc.add_table(rows=1, cols=2)
        t2.rows[0].cells[0].text = "Ресурс"
        t2.rows[0].cells[1].text = "Минут"
        for _, r in top_df.iterrows():
            row = t2.add_row().cells
            row[0].text = str(r["Resource"])
            row[1].text = str(int(r["Minutes"]))

    doc.save(str(out_path))
